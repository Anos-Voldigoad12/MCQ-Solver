#IMPORTS
import openai
import textract
import re
import json
import time
import docx
#SETUP
def setup():
    with open("api-key.txt",'r') as file:
        api_key = file.read().strip()
        if(len(api_key)==0):
            print("Key Not Found!")
            exit()
        else:
            openai.api_key = api_key        
#COMPLETION
def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        n=1,
        stop=None,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]
#SOLVING
def solveDocx(docxName):
    text = str(textract.process("./unsolved/"+docxName))
    questions = re.findall(r'"(.+?)"',text)
    questions = [ x.strip() for x in questions ]
    responses = list()
    try:
        for i in range(0,len(questions),10):
            prompt = f"""
            You will be provided with a list of questions.\
            Find out the correct answer to each question\
            (Each element is a string containing Correct option\
            along with the option value)\
            and output them in a python list.\
            [Do not include the question]

            {questions[i:i+10]}
            """
            response = get_completion(prompt)
            result = eval(response.strip())
            temp = dict()
            for x in range(len(result)):
                temp['question'] = eval("f'{}'".format(questions[i+x].strip()))
                temp['correct_answer'] = result[x]
                responses.append(temp.copy())
            print(f"Set-%d completed!"%((i/10)+1))
            time.sleep(30)
    finally:
        with open('./temp/'+docxName.split('.')[0]+'.json', 'w') as fout:
            json.dump(responses , fout)
#CONVERTING TO DOCX
def ouputToDocx(temp,outputName):
    f = open('./temp/'+temp)
    data = json.load(f)
    doc = docx.Document() #Create instance of Word Document
    heading = input("Heading: ")
    doc.add_heading(heading,0)
    for x in data:
        if("correct_answer" not in x.keys()):
            continue
        para = doc.add_paragraph(x["question"])
        para.add_run("\nCorrect Option : "+x["correct_answer"]).bold = True
    doc.save('./solved/'+outputName+".docx")


#DRIVER
setup()
docxName = input("Input File Name(with extension): ")
solveDocx(docxName)
outputName = input("Output File Name(without extension): ")
ouputToDocx(docxName.split('.')[0]+'.json',outputName)


    
